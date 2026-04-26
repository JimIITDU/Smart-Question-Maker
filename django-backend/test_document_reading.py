#!/usr/bin/env python
"""
Document Upload and Content Reading Test
Uploads PDF, DOCX, image, and text files, then reads their extracted content
"""
import requests
import json
import time
import uuid
from pathlib import Path
from docx import Document as DocxDocument
from PIL import Image, ImageDraw
import subprocess
import sys

BASE_URL = "http://localhost:8000/api"
TEST_FILES_DIR = Path(__file__).parent / "test_files"
UNIQUE_ID = uuid.uuid4().hex[:10]

# Test user
TEST_USER = {
    "username": f"testuser{UNIQUE_ID}",
    "email": f"testuser{UNIQUE_ID}@test.com",
    "password": "TestPass@123",
    "password2": "TestPass@123"
}

print("="*80)
print("DOCUMENT UPLOAD & CONTENT READING TEST")
print("="*80)

session = requests.Session()

# Step 1: Create test files
print("\n[1/5] Creating test files...")
TEST_FILES_DIR.mkdir(exist_ok=True)

# Create PDF using reportlab if available, else create image
pdf_path = TEST_FILES_DIR / "document.pdf"
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    c = canvas.Canvas(str(pdf_path), pagesize=letter)
    c.drawString(100, 750, "Introduction to Web Development")
    c.drawString(100, 730, "")
    c.drawString(100, 710, "Web development encompasses the work involved in developing websites for the Internet or an intranet.")
    c.drawString(100, 690, "It can range from developing a simple single static page of plain text to complex web applications,")
    c.drawString(100, 670, "electronic businesses, and social network services.")
    c.drawString(100, 650, "")
    c.drawString(100, 630, "Key Technologies:")
    c.drawString(100, 610, "- HTML5: Structure and semantics")
    c.drawString(100, 590, "- CSS3: Styling and layouts")
    c.drawString(100, 570, "- JavaScript: Interactivity and behavior")
    c.drawString(100, 550, "- Django/Flask: Backend frameworks")
    c.drawString(100, 530, "- PostgreSQL: Database management")
    c.save()
    print(f"  ✓ Created PDF: {pdf_path.name}")
except:
    print(f"  ~ PDF creation skipped (reportlab not installed)")
    pdf_path = None

# Create DOCX
docx_path = TEST_FILES_DIR / "document.docx"
doc = DocxDocument()
doc.add_heading("Cloud Computing Fundamentals", 0)
doc.add_paragraph(
    "Cloud computing is the on-demand availability of computer system resources, "
    "especially data storage and computing power, without direct active management by the user."
)
doc.add_heading("Benefits of Cloud Computing", 1)
for benefit in [
    "Cost Efficiency - Pay only for what you use",
    "Scalability - Easily scale resources up or down",
    "Reliability - Built-in redundancy and backups",
    "Security - Enterprise-grade security measures",
    "Flexibility - Access from anywhere with internet"
]:
    doc.add_paragraph(benefit, style='List Bullet')
doc.save(docx_path)
print(f"  ✓ Created DOCX: {docx_path.name}")

# Create text file
txt_path = TEST_FILES_DIR / "document.txt"
txt_path.write_text("""# Artificial Intelligence and Machine Learning

Artificial Intelligence (AI) is the simulation of human intelligence processes by computer systems.
These processes include learning, reasoning, and self-correction.

## Types of AI

1. Narrow AI (Weak AI): Designed to perform a specific task
   - Image recognition
   - Natural language processing
   - Recommendation systems

2. General AI (Strong AI): Hypothetical AI with human-level intelligence
   - Still largely theoretical
   - Active area of research

## Machine Learning Basics

Machine learning is a subset of AI where systems learn from data without explicit programming.

### Supervised Learning
- Labeled training data
- Examples: Classification, Regression

### Unsupervised Learning  
- Unlabeled data
- Examples: Clustering, Dimensionality reduction

### Reinforcement Learning
- Learn through interaction and feedback
- Examples: Game playing, Robotics
""")
print(f"  ✓ Created TXT: {txt_path.name}")

# Create image with text
img_path = TEST_FILES_DIR / "document.png"
img = Image.new('RGB', (800, 600), color='white')
draw = ImageDraw.Draw(img)
image_text = [
    "Data Science Fundamentals",
    "",
    "Data science combines statistics, programming, and domain expertise",
    "to extract meaningful insights from data.",
    "",
    "Key Components:",
    "1. Data Collection - Gather relevant data",
    "2. Data Cleaning - Handle missing/invalid data",
    "3. Exploratory Analysis - Understand data patterns",
    "4. Modeling - Build predictive models",
    "5. Visualization - Present findings"
]
y = 50
for line in image_text:
    draw.text((40, y), line, fill='black')
    y += 40
img.save(img_path)
print(f"  ✓ Created Image: {img_path.name}")

# Step 2: Register user
print("\n[2/5] Registering user...")
resp = session.post(
    f"{BASE_URL}/users/register/",
    json=TEST_USER,
    headers={"Accept": "application/json"}
)
if resp.status_code == 201:
    user_data = resp.json()
    print(f"  ✓ User registered: {TEST_USER['email']}")
else:
    print(f"  ✗ Failed: {resp.status_code} - {resp.json()}")
    sys.exit(1)

# Step 3: Get JWT token
print("\n[3/5] Getting JWT token...")
resp = session.post(
    f"{BASE_URL}/token/",
    json={"email": TEST_USER['email'], "password": TEST_USER['password']},
    headers={"Accept": "application/json"}
)
if resp.status_code == 200:
    token_data = resp.json()
    access_token = token_data['access']
    headers = {"Authorization": f"Bearer {access_token}"}
    print(f"  ✓ Token obtained")
else:
    print(f"  ✗ Failed: {resp.status_code}")
    sys.exit(1)

# Step 4: Upload documents
print("\n[4/5] Uploading documents...")
upload_results = []

test_files = [
    (docx_path, "Word Document (.docx)"),
    (txt_path, "Text File (.txt)"),
    (img_path, "Image File (.png)"),
]

# Add PDF if it was created
if pdf_path and pdf_path.exists():
    test_files.insert(0, (pdf_path, "PDF File (.pdf)"))

for file_path, file_type in test_files:
    print(f"\n  Uploading {file_type}...")
    with open(file_path, 'rb') as f:
        files = {'file': (file_path.name, f)}
        resp = session.post(f"{BASE_URL}/documents/", headers=headers, files=files)
    
    if resp.status_code == 201:
        doc_data = resp.json()
        doc_id = doc_data['id']
        upload_results.append({
            'file_type': file_type,
            'doc_id': doc_id,
            'status': doc_data.get('status'),
            'title': doc_data.get('title'),
        })
        print(f"    ✓ Uploaded successfully")
        print(f"    Document ID: {doc_id}")
        print(f"    Status: {doc_data.get('status')}")
    else:
        print(f"    ✗ Upload failed: {resp.status_code}")

# Step 5: Read extracted content
print("\n[5/5] Reading extracted content...")
print()

for result in upload_results:
    doc_id = result['doc_id']
    file_type = result['file_type']
    
    print(f"\n{'='*80}")
    print(f"Document: {result['title']} ({file_type})")
    print(f"ID: {doc_id} | Status: {result['status']}")
    print(f"{'='*80}")
    
    # Get document details
    resp = session.get(f"{BASE_URL}/documents/{doc_id}/", headers=headers)
    if resp.status_code == 200:
        doc_data = resp.json()
        
        # Display document metadata
        print(f"\nMetadata:")
        print(f"  - File Type: {doc_data.get('file_type')}")
        print(f"  - Status: {doc_data.get('status')}")
        print(f"  - Pages: {doc_data.get('page_count', 'N/A')}")
        print(f"  - Created: {doc_data.get('created_at')}")
        
        # Display extracted content
        extracted_text = doc_data.get('extracted_text', '')
        if extracted_text:
            print(f"\nExtracted Content ({len(extracted_text)} characters):")
            print("-" * 80)
            # Show first 500 characters or full content if shorter
            if len(extracted_text) > 500:
                print(extracted_text[:500])
                print(f"\n... (truncated, total: {len(extracted_text)} characters)")
            else:
                print(extracted_text)
            print("-" * 80)
        else:
            print(f"\n[No extracted content yet - Status: {doc_data.get('status')}]")
    else:
        print(f"  ✗ Failed to get document details: {resp.status_code}")

print(f"\n{'='*80}")
print("Test Complete!")
print(f"{'='*80}")
print(f"\nTest Summary:")
print(f"  - Total files uploaded: {len(upload_results)}")
print(f"  - Test files location: {TEST_FILES_DIR}")
print(f"\nAPI Endpoints tested:")
print(f"  - POST /api/users/register/ - User registration")
print(f"  - POST /api/token/ - Get JWT token")
print(f"  - POST /api/documents/ - Upload document")
print(f"  - GET  /api/documents/{{id}}/ - Get document details with extracted content")
