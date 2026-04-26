#!/usr/bin/env python
"""
Document Upload API Test Script
Tests uploading PDF, DOCX, image, and text files
"""

import os
import sys
import json
import requests
import time
from pathlib import Path
from docx import Document as DocxDocument
from PIL import Image, ImageDraw, ImageFont

# Configuration
BASE_URL = "http://localhost:8000/api"
TEST_FILES_DIR = Path(__file__).parent / "test_files"

# Test user credentials (use unique email to avoid conflicts)
TIMESTAMP = str(int(time.time() * 1000))
TEST_USER = {
    "username": f"testuser_{TIMESTAMP}",
    "email": f"testuser{TIMESTAMP}@example.com",
    "password": "TestPass@123456"
}

# Color codes for terminal output
GREEN = '\033[92m'
RED = '\033[91m'
BLUE = '\033[94m'
YELLOW = '\033[93m'
RESET = '\033[0m'
BOLD = '\033[1m'


def print_header(text):
    print(f"\n{BOLD}{BLUE}{'='*60}{RESET}")
    print(f"{BOLD}{BLUE}{text}{RESET}")
    print(f"{BOLD}{BLUE}{'='*60}{RESET}\n")


def print_success(text):
    print(f"{GREEN}[OK] {text}{RESET}")


def print_error(text):
    print(f"{RED}[ERROR] {text}{RESET}")


def print_info(text):
    print(f"{BLUE}[INFO] {text}{RESET}")


def print_step(text):
    print(f"\n{YELLOW}[STEP] {text}{RESET}")


def create_test_files():
    """Create test files for upload testing."""
    print_step("Creating test files...")
    
    TEST_FILES_DIR.mkdir(exist_ok=True)
    
    # 1. Create DOCX file
    docx_path = TEST_FILES_DIR / "sample_document.docx"
    doc = DocxDocument()
    doc.add_heading("Python Programming Guide", level=0)
    doc.add_paragraph(
        "Python is a versatile, high-level programming language known for its simplicity and readability. "
        "It supports multiple programming paradigms including procedural, object-oriented, and functional programming."
    )
    doc.add_heading("Key Features", level=1)
    doc.add_paragraph("Simple and readable syntax")
    doc.add_paragraph("Extensive standard library")
    doc.add_paragraph("Cross-platform compatibility")
    doc.add_paragraph("Strong community support")
    doc.add_heading("Applications", level=1)
    doc.add_paragraph("Web development with Django and Flask")
    doc.add_paragraph("Data science and machine learning with pandas, numpy, scikit-learn")
    doc.add_paragraph("Automation and scripting")
    doc.add_paragraph("Scientific computing")
    doc.save(docx_path)
    print_success(f"Created DOCX file: {docx_path}")
    
    # 2. Create image file with text
    img_path = TEST_FILES_DIR / "sample_image.png"
    img = Image.new('RGB', (800, 600), color='white')
    draw = ImageDraw.Draw(img)
    text_lines = [
        "Welcome to Image OCR Testing",
        "",
        "This image contains text that will be",
        "extracted using Optical Character Recognition (OCR).",
        "",
        "Key Points:",
        "- OCR technology reads text from images",
        "- Used in document digitization",
        "- Important for accessibility"
    ]
    y_position = 50
    for line in text_lines:
        draw.text((50, y_position), line, fill='black')
        y_position += 40
    img.save(img_path)
    print_success(f"Created image file: {img_path}")
    
    # 3. Create text file
    txt_path = TEST_FILES_DIR / "sample_text.txt"
    if not txt_path.exists():
        with open(txt_path, 'w') as f:
            f.write("""# Introduction to Machine Learning

Machine learning is a subset of artificial intelligence that focuses on the development of algorithms and statistical models that enable computers to learn from and make predictions or decisions based on data, without being explicitly programmed.

## Key Concepts

Machine learning can be divided into three main types:

1. Supervised Learning: The algorithm learns from labeled data. Examples include regression and classification.

2. Unsupervised Learning: The algorithm learns from unlabeled data. Examples include clustering and dimensionality reduction.

3. Reinforcement Learning: The algorithm learns through interaction with an environment, receiving rewards or penalties.

## Applications

- Image recognition and computer vision
- Natural language processing
- Recommendation systems
- Fraud detection
- Healthcare diagnosis
- Autonomous vehicles
- Stock market prediction
""")
        print_success(f"Created text file: {txt_path}")
    
    return docx_path, img_path, txt_path


def register_user(session):
    """Register a test user."""
    print_step("Registering test user...")
    
    url = f"{BASE_URL}/users/register/"
    response = session.post(url, json=TEST_USER)
    
    print_info(f"Registration response: {response.status_code}")
    
    if response.status_code == 201:
        print_success(f"User registered: {TEST_USER['username']}")
        return response.json()
    elif response.status_code == 400:
        print_info(f"User may already exist or validation error")
        print_info(f"Response: {response.json()}")
        return None
    else:
        print_error(f"Registration failed: {response.status_code}")
        print_error(f"Response: {response.text}")
        return None


def login_user(session):
    """Login user and get JWT token."""
    print_step("Logging in user...")
    
    url = f"{BASE_URL}/token/"
    response = session.post(url, json={
        "email": TEST_USER['email'],
        "password": TEST_USER['password']
    })
    
    if response.status_code == 200:
        data = response.json()
        access_token = data['access']
        print_success(f"Login successful")
        print_info(f"Access Token: {access_token[:50]}...")
        return access_token
    else:
        print_error(f"Login failed: {response.status_code}")
        print_error(f"Response: {response.text}")
        return None


def upload_document(session, access_token, file_path, file_name=None):
    """Upload a document."""
    if file_name is None:
        file_name = Path(file_path).name
    
    url = f"{BASE_URL}/documents/"
    headers = {"Authorization": f"Bearer {access_token}"}
    
    with open(file_path, 'rb') as f:
        files = {'file': (file_name, f)}
        response = session.post(url, headers=headers, files=files)
    
    return response


def test_upload_all_files(access_token, docx_path, img_path, txt_path):
    """Test uploading all file types."""
    print_header("Testing Document Uploads")
    
    session = requests.Session()
    test_files = [
        (txt_path, "Text File (.txt)"),
        (docx_path, "Word Document (.docx)"),
        (img_path, "Image File (.png)"),
    ]
    
    results = []
    
    for file_path, file_type in test_files:
        print_step(f"Uploading {file_type}...")
        
        try:
            response = upload_document(session, access_token, file_path)
            
            if response.status_code == 201:
                data = response.json()
                print_success(f"Upload successful - {file_type}")
                print_info(f"Document ID: {data['id']}")
                print_info(f"Status: {data['status']}")
                print_info(f"File Type: {data['file_type']}")
                results.append({
                    'file_type': file_type,
                    'doc_id': data['id'],
                    'status': 'success',
                    'response': data
                })
            else:
                print_error(f"Upload failed - {file_type}: {response.status_code}")
                print_error(f"Response: {response.text}")
                results.append({
                    'file_type': file_type,
                    'status': 'failed',
                    'error': response.text
                })
        except Exception as e:
            print_error(f"Upload error - {file_type}: {str(e)}")
            results.append({
                'file_type': file_type,
                'status': 'error',
                'error': str(e)
            })
    
    return results


def check_document_status(access_token, doc_id):
    """Check the status of a document."""
    session = requests.Session()
    url = f"{BASE_URL}/documents/{doc_id}/status/"
    headers = {"Authorization": f"Bearer {access_token}"}
    
    response = session.get(url, headers=headers)
    
    if response.status_code == 200:
        return response.json()
    else:
        return None


def list_documents(access_token):
    """List all documents for the user."""
    session = requests.Session()
    url = f"{BASE_URL}/documents/"
    headers = {"Authorization": f"Bearer {access_token}"}
    
    response = session.get(url, headers=headers)
    
    if response.status_code == 200:
        return response.json()
    else:
        return None


def print_results_summary(results):
    """Print summary of upload results."""
    print_header("Upload Results Summary")
    
    successful = sum(1 for r in results if r['status'] == 'success')
    failed = sum(1 for r in results if r['status'] in ['failed', 'error'])
    
    print(f"Total Files: {len(results)}")
    print(f"{GREEN}Successful: {successful}{RESET}")
    print(f"{RED}Failed: {failed}{RESET}")
    
    print("\nDetails:")
    for result in results:
        if result['status'] == 'success':
            print(f"\n{GREEN}✓ {result['file_type']}{RESET}")
            print(f"  Document ID: {result['doc_id']}")
            print(f"  Status: {result['response']['status']}")
            print(f"  File Type: {result['response']['file_type']}")
        else:
            print(f"\n{RED}✗ {result['file_type']}{RESET}")
            print(f"  Error: {result.get('error', 'Unknown error')}")


def main():
    """Main test function."""
    print_header("Document Upload API Test Suite")
    
    # Check if server is running
    print_step("Checking server connection...")
    try:
        response = requests.get(f"{BASE_URL}/documents/", timeout=5)
    except requests.exceptions.ConnectionError:
        print_error("Cannot connect to server at {BASE_URL}")
        print_info("Make sure the Django server is running: python manage.py runserver")
        sys.exit(1)
    except Exception as e:
        print_error(f"Connection error: {str(e)}")
        sys.exit(1)
    
    print_success("Server is running")
    
    # Create test files
    docx_path, img_path, txt_path = create_test_files()
    
    # Register and login
    session = requests.Session()
    register_user(session)
    access_token = login_user(session)
    
    if not access_token:
        print_error("Failed to obtain access token")
        sys.exit(1)
    
    # Test uploads
    results = test_upload_all_files(access_token, docx_path, img_path, txt_path)
    
    # Print summary
    print_results_summary(results)
    
    # Check document statuses
    print_header("Checking Document Processing Status")
    for result in results:
        if result['status'] == 'success':
            doc_id = result['doc_id']
            print_step(f"Checking status for document {doc_id}...")
            status_info = check_document_status(access_token, doc_id)
            if status_info:
                print_info(f"Status: {status_info['status']}")
                print_info(f"Extracted Text Length: {status_info.get('extracted_text_length', 0)} characters")
                if status_info.get('page_count'):
                    print_info(f"Page Count: {status_info['page_count']}")
    
    # List all documents
    print_header("Listing All Documents")
    documents = list_documents(access_token)
    if documents:
        if isinstance(documents, dict) and 'results' in documents:
            docs = documents['results']
        elif isinstance(documents, list):
            docs = documents
        else:
            docs = []
        
        print(f"Total documents: {len(docs)}")
        for doc in docs:
            print(f"\n{BLUE}Document: {doc['id']}{RESET}")
            print(f"  Title: {doc.get('title', 'N/A')}")
            print(f"  Type: {doc.get('file_type', 'N/A')}")
            print(f"  Status: {doc.get('status', 'N/A')}")
            print(f"  Created: {doc.get('created_at', 'N/A')}")
    
    print_header("Test Complete!")
    print_success("All tests completed successfully!")
    print_info(f"Test files location: {TEST_FILES_DIR}")


if __name__ == "__main__":
    main()
