#!/usr/bin/env python
"""
Simple Document Upload Test
Tests uploading different file types with proper authentication
"""
import requests
import json
import time
import uuid
from pathlib import Path
from docx import Document as DocxDocument
from PIL import Image, ImageDraw

BASE_URL = "http://localhost:8000/api"
TEST_FILES_DIR = Path(__file__).parent / "test_files"
UNIQUE_ID = str(uuid.uuid4())[:8]

# Create test user
TEST_USER = {
    "username": f"user{UNIQUE_ID}",
    "email": f"user{UNIQUE_ID}@test.com",
    "password": "SecurePass@123"
}

print("="*70)
print("Document Upload API - Comprehensive Test")
print("="*70)

session = requests.Session()

# Step 1: Create test files
print("\n[SETUP] Creating test files...")
TEST_FILES_DIR.mkdir(exist_ok=True)

docx_path = TEST_FILES_DIR / "document.docx"
doc = DocxDocument()
doc.add_heading("Python Fundamentals", 0)
doc.add_paragraph("Python is a versatile, interpreted programming language.")
doc.add_heading("Data Types", 1)
for dtype in ["Strings", "Integers", "Lists", "Dictionaries", "Tuples"]:
    doc.add_paragraph(dtype, style='List Bullet')
doc.save(docx_path)
print(f"  ✓ Created DOCX: {docx_path.name}")

txt_path = TEST_FILES_DIR / "notes.txt"
txt_path.write_text("""# Web Development Guide

## Frontend Technologies
- HTML5 for structure
- CSS3 for styling  
- JavaScript for interactivity
- React or Vue.js frameworks

## Backend Technologies
- Django or FastAPI
- NodeJS with Express
- Python with Flask
- Database: PostgreSQL or MongoDB

## Full Stack Development
Combines both frontend and backend technologies
to create complete web applications.
""")
print(f"  ✓ Created TXT: {txt_path.name}")

img_path = TEST_FILES_DIR / "diagram.png"
img = Image.new('RGB', (600, 400), color='white')
draw = ImageDraw.Draw(img)
lines = [
    "Architecture Diagram",
    "Frontend - Backend - Database",
    "",
    "HTTP Requests flow between layers",
    "RESTful API for communication",
    "Microservices approach recommended"
]
y = 50
for line in lines:
    draw.text((40, y), line, fill='black')
    y += 50
img.save(img_path)
print(f"  ✓ Created PNG: {img_path.name}")

# Step 2: Register user
print("\n[AUTH] Registering user...")
resp = session.post(
    f"{BASE_URL}/users/register/", 
    json=TEST_USER,
    headers={"Accept": "application/json"}
)
if resp.status_code == 201:
    print(f"  ✓ User registered: {TEST_USER['email']}")
else:
    print(f"  ✗ Registration failed: {resp.status_code}")
    try:
        print(f"    {resp.json()}")
    except:
        print(f"    Response: {resp.text[:300]}")
    exit(1)

# Step 3: Get JWT token
print("\n[AUTH] Obtaining JWT token...")
token_resp = session.post(f"{BASE_URL}/token/", json={
    "email": TEST_USER['email'],
    "password": TEST_USER['password']
})
if token_resp.status_code == 200:
    token_data = token_resp.json()
    access_token = token_data['access']
    print(f"  ✓ Token obtained")
    print(f"    Token: {access_token[:40]}...")
else:
    print(f"  ✗ Token request failed: {token_resp.status_code}")
    print(f"    {token_resp.text[:200]}")
    exit(1)

# Step 4: Test document uploads
print("\n[UPLOAD] Testing document uploads...")
headers = {"Authorization": f"Bearer {access_token}"}

test_files = [
    (txt_path, "Text File (.txt)"),
    (docx_path, "Word Document (.docx)"),
    (img_path, "Image File (.png)"),
]

upload_results = []

for file_path, file_type in test_files:
    print(f"\n  Uploading {file_type}...")
    with open(file_path, 'rb') as f:
        files = {'file': (file_path.name, f)}
        resp = session.post(f"{BASE_URL}/documents/", headers=headers, files=files)
    
    if resp.status_code == 201:
        data = resp.json()
        doc_id = data['id']
        upload_results.append((file_type, doc_id, 'success'))
        print(f"    ✓ Upload successful")
        print(f"      Document ID: {doc_id}")
        print(f"      Status: {data.get('status', 'unknown')}")
        print(f"      File Type: {data.get('file_type', 'unknown')}")
    else:
        print(f"    ✗ Upload failed: {resp.status_code}")
        print(f"      Response: {resp.json()}")
        upload_results.append((file_type, None, 'failed'))

# Step 5: List documents
print("\n[LIST] Fetching document list...")
resp = session.get(f"{BASE_URL}/documents/", headers=headers)
if resp.status_code == 200:
    data = resp.json()
    docs = data.get('results', data) if isinstance(data, dict) else data
    print(f"  ✓ Retrieved {len(docs)} document(s)")
    for doc in docs[:5]:  # Show first 5
        print(f"    - {doc.get('title', 'Untitled')} (ID: {doc.get('id')}, Status: {doc.get('status')})")
else:
    print(f"  ✗ Failed to list documents: {resp.status_code}")

# Step 6: Check document status
print("\n[STATUS] Checking document processing status...")
for file_type, doc_id, status in upload_results:
    if status == 'success' and doc_id:
        resp = session.get(f"{BASE_URL}/documents/{doc_id}/status/", headers=headers)
        if resp.status_code == 200:
            status_data = resp.json()
            print(f"  {file_type}:")
            print(f"    - Status: {status_data.get('status')}")
            print(f"    - Extracted text length: {status_data.get('extracted_text_length', 0)} chars")
            if status_data.get('page_count'):
                print(f"    - Pages: {status_data.get('page_count')}")
        else:
            print(f"  {file_type}: Failed to get status ({resp.status_code})")

# Step 7: Summary
print("\n" + "="*70)
print("SUMMARY")
print("="*70)
success_count = sum(1 for _, _, s in upload_results if s == 'success')
print(f"Total uploads: {len(upload_results)}")
print(f"Successful: {success_count}")
print(f"Failed: {len(upload_results) - success_count}")

print("\nTest files location:")
print(f"  {TEST_FILES_DIR}")

print("\n" + "="*70)
if success_count == len(upload_results):
    print("All tests passed! ✓")
else:
    print("Some tests failed. Check errors above.")
print("="*70)
